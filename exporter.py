import re
from datetime import datetime
from pathlib import Path


def _get_filename(ext, title="회의록", save_folder=None):
    from config import get_save_folder
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r'[\\/:*?"<>|]', "_", title)
    folder = save_folder or get_save_folder()
    return str(Path(folder) / f"{safe_title}_{timestamp}.{ext}")


def _parse_transcript_lines(full_text):
    """전사 텍스트를 (타임스탬프, 화자, 내용) 튜플 리스트로 파싱."""
    pattern = re.compile(r"\[(\d{2}:\d{2})\]\s*(화자\s*\w+):\s*(.+)")
    parsed = []
    for line in full_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        m = pattern.match(line)
        if m:
            parsed.append((m.group(1), m.group(2), m.group(3)))
        else:
            parsed.append(("", "", line))
    return parsed


# ═══════════════════════════════════════════════════════════════════
# Excel 내보내기
# scope: "full" (전체), "summary_only" (요약만), "transcript_only" (기록만)
# ═══════════════════════════════════════════════════════════════════
def save_excel(full_text, summary, metadata=None, scope="full"):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = Workbook()
    header_fill = PatternFill(start_color="2196F3", end_color="2196F3", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11, name="맑은 고딕")
    body_font = Font(size=10, name="맑은 고딕")
    title_font = Font(size=14, bold=True, name="맑은 고딕")
    thin_border = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )
    alt_fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")

    title = metadata.title if metadata else "회의록"

    # ── 시트 1: 회의 정보 (항상 포함) ──
    ws_info = wb.active
    ws_info.title = "회의 정보"
    ws_info.column_dimensions["A"].width = 18
    ws_info.column_dimensions["B"].width = 50

    ws_info["A1"] = title
    ws_info["A1"].font = title_font
    ws_info.merge_cells("A1:B1")

    info_items = [
        ("날짜", metadata.date if metadata else datetime.now().strftime("%Y-%m-%d")),
        ("참석자", metadata.attendees if metadata and metadata.attendees else "-"),
        ("회의 유형", metadata.meeting_type if metadata else "일반 회의"),
        ("장소", metadata.location if metadata and metadata.location else "-"),
        ("생성 시간", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    ]
    for i, (key, val) in enumerate(info_items, start=3):
        cell_a = ws_info[f"A{i}"]
        cell_b = ws_info[f"B{i}"]
        cell_a.value = key
        cell_a.font = Font(bold=True, size=10, name="맑은 고딕")
        cell_a.fill = PatternFill(start_color="E3F2FD", end_color="E3F2FD", fill_type="solid")
        cell_a.border = thin_border
        cell_b.value = val
        cell_b.font = body_font
        cell_b.border = thin_border

    # ── 요약 시트 ──
    if scope in ("full", "summary_only"):
        ws_summary = wb.create_sheet("요약")
        ws_summary.column_dimensions["A"].width = 20
        ws_summary.column_dimensions["B"].width = 70
        ws_summary["A1"] = "회의 요약"
        ws_summary["A1"].font = title_font
        ws_summary.merge_cells("A1:B1")
        for i, line in enumerate(summary.split("\n"), start=3):
            line = line.strip()
            if not line:
                continue
            cell = ws_summary[f"A{i}"]
            if line.startswith("##"):
                cell.value = line.replace("##", "").strip()
                cell.font = Font(bold=True, size=11, name="맑은 고딕", color="1565C0")
                ws_summary.merge_cells(f"A{i}:B{i}")
            else:
                cell.value = line
                cell.font = body_font

    # ── 전체 기록 시트 ──
    if scope in ("full", "transcript_only"):
        ws_full = wb.create_sheet("전체 기록")
        headers = ["타임스탬프", "화자", "내용"]
        col_widths = [14, 14, 80]
        for col_idx, (header, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws_full.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")
            cell.border = thin_border
            ws_full.column_dimensions[chr(64 + col_idx)].width = width
        ws_full.auto_filter.ref = "A1:C1"
        ws_full.freeze_panes = "A2"
        parsed = _parse_transcript_lines(full_text)
        for row_idx, (ts, speaker, content) in enumerate(parsed, start=2):
            ws_full.cell(row=row_idx, column=1, value=ts).font = body_font
            ws_full.cell(row=row_idx, column=2, value=speaker).font = body_font
            ws_full.cell(row=row_idx, column=3, value=content).font = body_font
            for col in range(1, 4):
                cell = ws_full.cell(row=row_idx, column=col)
                cell.border = thin_border
                if row_idx % 2 == 0:
                    cell.fill = alt_fill

    path = _get_filename("xlsx", title)
    wb.save(path)
    return path


# ═══════════════════════════════════════════════════════════════════
# Word 내보내기
# ═══════════════════════════════════════════════════════════════════
def save_word(full_text, summary, metadata=None, scope="full"):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    title = metadata.title if metadata else "회의록"

    # ── 페이지 설정 ──
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 바닥글: 페이지 번호 ──
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._element.append(fldChar1)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run._element.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar2)

    # ── 제목 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = Pt(6)
    run = title_para.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    # 메타데이터 테이블
    date_str = metadata.date if metadata else datetime.now().strftime("%Y-%m-%d")
    meta_items = [("날짜", date_str)]
    if metadata and metadata.attendees:
        meta_items.append(("참석자", metadata.attendees))
    if metadata and metadata.meeting_type:
        meta_items.append(("회의 유형", metadata.meeting_type))
    if metadata and metadata.location:
        meta_items.append(("장소", metadata.location))

    table = doc.add_table(rows=len(meta_items), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (key, val) in enumerate(meta_items):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for r in paragraph.runs:
                    r.font.size = Pt(10)
                    r.font.name = "맑은 고딕"

    doc.add_paragraph()

    def add_divider():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("─" * 50)
        r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        r.font.size = Pt(8)

    # ── 요약 섹션 ──
    if scope in ("full", "summary_only"):
        doc.add_heading("요약", level=1)
        for line in summary.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("- [ ] ") or line.startswith("- [x] "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
                for r in p.runs:
                    r.font.size = Pt(10)
            elif line.startswith("- "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
                for r in p.runs:
                    r.font.size = Pt(10)
            elif line[0:1].isdigit() and "." in line[:4]:
                p = doc.add_paragraph(line, style="List Number")
                for r in p.runs:
                    r.font.size = Pt(10)
            else:
                p = doc.add_paragraph(line)
                for r in p.runs:
                    r.font.size = Pt(10)

    # ── 전체 기록 섹션 ──
    if scope in ("full", "transcript_only"):
        if scope == "full":
            add_divider()
        doc.add_heading("전체 기록", level=1)
        parsed = _parse_transcript_lines(full_text)
        for ts, speaker, content in parsed:
            p = doc.add_paragraph()
            if ts and speaker:
                run_ts = p.add_run(f"[{ts}] ")
                run_ts.font.size = Pt(9)
                run_ts.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
                run_sp = p.add_run(f"{speaker}: ")
                run_sp.font.size = Pt(10)
                run_sp.font.bold = True
                run_sp.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
                run_ct = p.add_run(content)
                run_ct.font.size = Pt(10)
            else:
                r = p.add_run(content)
                r.font.size = Pt(10)
                if content.startswith("[세션"):
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

    path = _get_filename("docx", title)
    doc.save(path)
    return path


# ═══════════════════════════════════════════════════════════════════
# PDF 내보내기
# ═══════════════════════════════════════════════════════════════════
def save_pdf(full_text, summary, metadata=None, scope="full"):
    from fpdf import FPDF
    import os

    title = metadata.title if metadata else "회의록"
    date_str = metadata.date if metadata else datetime.now().strftime("%Y-%m-%d")

    class MeetingPDF(FPDF):
        def header(self):
            self.set_font("korean", size=8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 8, "회의록", align="L")
            self.cell(0, 8, date_str, align="R", new_x="LMARGIN", new_y="NEXT")
            self.set_draw_color(200, 200, 200)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font("korean", size=8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, f"- {self.page_no()}/{{nb}} -", align="C")

    pdf = MeetingPDF()
    pdf.alias_nb_pages()

    font_regular = "C:/Windows/Fonts/malgun.ttf"
    font_bold = "C:/Windows/Fonts/malgunbd.ttf"
    if not os.path.exists(font_regular):
        font_regular = "C:/Windows/Fonts/gulim.ttc"
        font_bold = font_regular

    pdf.add_font("korean", "", font_regular, uni=True)
    if os.path.exists(font_bold) and font_bold != font_regular:
        pdf.add_font("korean_bold", "", font_bold, uni=True)
    else:
        pdf.add_font("korean_bold", "", font_regular, uni=True)

    pdf.set_auto_page_break(auto=True, margin=20)

    # ── 제목 페이지 ──
    pdf.add_page()
    pdf.ln(30)
    pdf.set_font("korean_bold", size=22)
    pdf.set_text_color(21, 101, 192)
    pdf.cell(0, 15, title, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(8)

    pdf.set_font("korean", size=10)
    pdf.set_text_color(66, 66, 66)
    meta_lines = [f"날짜: {date_str}"]
    if metadata and metadata.attendees:
        meta_lines.append(f"참석자: {metadata.attendees}")
    if metadata and metadata.meeting_type:
        meta_lines.append(f"회의 유형: {metadata.meeting_type}")
    if metadata and metadata.location:
        meta_lines.append(f"장소: {metadata.location}")
    for ml in meta_lines:
        pdf.cell(0, 7, ml, align="C", new_x="LMARGIN", new_y="NEXT")

    # ── 요약 섹션 ──
    if scope in ("full", "summary_only"):
        pdf.add_page()
        pdf.set_draw_color(33, 150, 243)
        pdf.set_fill_color(33, 150, 243)
        pdf.rect(10, pdf.get_y(), 190, 0.8, "F")
        pdf.ln(4)

        for line in summary.split("\n"):
            line = line.strip()
            if not line:
                pdf.ln(3)
                continue
            if line.startswith("## "):
                pdf.ln(4)
                pdf.set_font("korean_bold", size=13)
                pdf.set_text_color(21, 101, 192)
                pdf.cell(0, 9, line[3:], new_x="LMARGIN", new_y="NEXT")
            elif line.startswith("[ ") or line.startswith("["):
                pdf.set_font("korean_bold", size=14)
                pdf.set_text_color(21, 101, 192)
                pdf.cell(0, 10, line, new_x="LMARGIN", new_y="NEXT")
            else:
                pdf.set_font("korean", size=10)
                pdf.set_text_color(33, 33, 33)
                pdf.multi_cell(0, 6.5, line)

    # ── 전체 기록 섹션 ──
    if scope in ("full", "transcript_only"):
        pdf.add_page()
        pdf.set_font("korean_bold", size=14)
        pdf.set_text_color(21, 101, 192)
        pdf.cell(0, 10, "전체 기록", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(33, 150, 243)
        pdf.rect(10, pdf.get_y(), 190, 0.8, "F")
        pdf.ln(4)

        parsed = _parse_transcript_lines(full_text)
        for ts, speaker, content in parsed:
            if ts and speaker:
                pdf.set_font("korean", size=8)
                pdf.set_text_color(117, 117, 117)
                pdf.cell(18, 6, f"[{ts}]")
                pdf.set_font("korean_bold", size=10)
                pdf.set_text_color(21, 101, 192)
                pdf.cell(22, 6, f"{speaker}:")
                pdf.set_font("korean", size=10)
                pdf.set_text_color(33, 33, 33)
                pdf.multi_cell(0, 6, content)
            else:
                pdf.set_font("korean", size=10)
                if content.startswith("[세션"):
                    pdf.ln(4)
                    pdf.set_font("korean_bold", size=11)
                    pdf.set_text_color(66, 66, 66)
                    pdf.cell(0, 8, content, new_x="LMARGIN", new_y="NEXT")
                    pdf.ln(2)
                else:
                    pdf.set_text_color(33, 33, 33)
                    pdf.multi_cell(0, 6, content)

    path = _get_filename("pdf", title)
    pdf.output(path)
    return path


# ═══════════════════════════════════════════════════════════════════
# Markdown 내보내기
# ═══════════════════════════════════════════════════════════════════
def save_markdown(full_text, summary, metadata=None, scope="full"):
    title = metadata.title if metadata else "회의록"
    date_str = metadata.date if metadata else datetime.now().strftime("%Y-%m-%d")

    lines = [f"# {title}\n"]

    # 메타데이터 테이블
    lines.append("| 항목 | 내용 |")
    lines.append("|------|------|")
    lines.append(f"| 날짜 | {date_str} |")
    if metadata and metadata.attendees:
        lines.append(f"| 참석자 | {metadata.attendees} |")
    if metadata and metadata.meeting_type:
        lines.append(f"| 회의 유형 | {metadata.meeting_type} |")
    if metadata and metadata.location:
        lines.append(f"| 장소 | {metadata.location} |")
    lines.append(f"| 생성 시간 | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |")
    lines.append("")

    # 요약 섹션
    if scope in ("full", "summary_only"):
        lines.append("---\n")
        lines.append("## 요약\n")
        for line in summary.split("\n"):
            line = line.strip()
            if not line:
                lines.append("")
                continue
            # ## 헤더는 ### 로 한 단계 낮춤
            if line.startswith("## "):
                lines.append(f"### {line[3:]}")
            elif line.startswith("[ ") or line.startswith("["):
                lines.append(f"**{line}**")
            else:
                lines.append(line)
        lines.append("")

    # 전체 기록 섹션
    if scope in ("full", "transcript_only"):
        lines.append("---\n")
        lines.append("## 전체 기록\n")
        parsed = _parse_transcript_lines(full_text)
        for ts, speaker, content in parsed:
            if ts and speaker:
                lines.append(f"> **[{ts}] {speaker}:** {content}")
            else:
                if content.startswith("[세션"):
                    lines.append(f"\n### {content}\n")
                else:
                    lines.append(content)
        lines.append("")

    md_content = "\n".join(lines)
    path = _get_filename("md", title)
    with open(path, "w", encoding="utf-8") as f:
        f.write(md_content)
    return path
